import time
import requests
import json
import re
import os
from datetime import datetime
from lxml import etree
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from PIL import Image

session = requests.Session()

def create_path(path: str) -> None:
    if not os.path.exists(path):
        os.makedirs(path)


def get_article_html(cookie: str, article_id: str):
    """
    get article html
    :param cookie: website's cookie information
    :param article_id: article's url such as 1097430290934005800 (https://www.bilibili.com/opus/1097430290934005800?spm_id_from=333.1387.0.0)
    :return: html
    """
    url = f"https://www.bilibili.com/opus/{article_id}"
    headers = {
        'cookie': cookie,
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36'
    }
    response = session.get(url, headers=headers)
    return response.text


def get_text(item):
    pass


def get_p_tag_text(html) -> tuple:
    result = html.xpath('./span | ./strong')
    color = result[0].get('style')
    if color is not None:
        color = re.findall(r'color:(.*?);', result[0].get('style'))[0]
    is_strong = result[0].tag == 'strong'
    if len(result) != 0:
        return ''.join(result[i].text for i in range(0, len(result))), color, is_strong
    else:
        return '\n', color, is_strong


def get_img(cookie, url_list: list, path) -> list:
    """
    get img.
    :param cookie: website's cookie information
    :param url_list: each item under the opus-module-content class
    :param path: img path
    :return: image path list
    """
    result = list()

    headers = {
        'cookie': cookie,
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36'
    }
    for url in url_list:
        response = session.get(url, headers=headers)
        time.sleep(0.5)

        create_path(path)

        name = url.split('/')[-1]

        with open(f'{path}/{name}', 'wb') as f:
            f.write(response.content)

        print(f'{path}/{name}')

        result.append(f'{path}/{name}')

    return result


def add_text(doc, item: dict) -> None:
    """
    Add text to document.
    :param doc: docx document
    :param item: item's text
    :return: None
    """
    is_first = True
    for nodes in item['nodes']:
        words = nodes['word']['words']
        if len(item['nodes']) > 1 and is_first is False:
            paragraph.add_run(words)
        else:
            paragraph = doc.add_paragraph(words)

        # get color
        if 'color' in nodes['word'].keys():
            color = nodes['word']['color'].strip('#')
            r = int(color[0:2], 16)
            g = int(color[2:4], 16)
            b = int(color[4:6], 16)
        else:
            r = 47
            g = 50
            b = 56

        fontSize = nodes['word']['font_size']

        # get bold
        if 'bold' in nodes['word']['style'].keys():
            is_blod = nodes['word']['style']['bold']
        else:
            is_blod = False

        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(r, g, b)
            run.font.size = Pt(fontSize / 1.5)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.bold = is_blod

        is_first = False


def add_image(doc, image_path_list) -> None:
    """
    Add img to document.
    :param doc: docx document
    :param image_path_list: image path list
    :return: None
    """
    section = doc.sections[0]

    available_width = (section.page_width - section.left_margin - section.right_margin)

    for image_path in image_path_list:
        img = Image.open(image_path)
        img.save(image_path)
        img.close()
        doc.add_picture(image_path, width=available_width)

def add_link(doc, html) -> tuple:
    """
    Add link to document.
    :param doc: docx document
    :param html: html.
    :return: link title,
    """


def get_module(INITIAL_STATE: dict, module_name: str):
    INITIAL_STATE = INITIAL_STATE['detail']['modules']
    for item in INITIAL_STATE:
        if module_name in ''.join(list(item.keys())):
            return item
    return None

def get_article(cookie: str, article_id: str, doc_storage_location: str = '.', document_name: str = 'Document.doc',
                img_path: str = 'img') -> None:
    """
    get_article
    :param cookie: website's cookie information.
    :param article_id: article's url such as 1097430290934005800 (https://www.bilibili.com/opus/1097430290934005800?spm_id_from=333.1387.0.0)
    :param doc_storage_location: doc local storage in word.
    :param img_path: img path
    :return: None
    """
    html = get_article_html(cookie, article_id)
    INITIAL_STATE = re.findall(r'window.__INITIAL_STATE__=(.*);\(function', html)[0]
    INITIAL_STATE = json.loads(INITIAL_STATE)

    with open('test.json', 'w', encoding='utf-8') as f:
        json.dump(INITIAL_STATE, f, ensure_ascii=False)

    element = etree.HTML(html)

    title = get_module(INITIAL_STATE, 'title')
    author = get_module(INITIAL_STATE, 'author')
    module_top = get_module(INITIAL_STATE, 'module_top')
    module_content = get_module(INITIAL_STATE, 'module_content')

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    if title is not None:
        title = title['module_title']['text']
        doc_title = doc.add_heading(title, level=1)
        for run in doc_title.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(22 / 1.5)

    if author is not None:
        author_name = author['module_author']['name']
        author_time = author['module_author']['pub_time']

        if '编辑于' in author_time:
            author_time = datetime.strptime(author_time, '编辑于 %Y年%m月%d日 %H:%M')
        else:
            author_time = datetime.strptime(author_time, '%Y年%m月%d日 %H:%M')
    else:
        author_name = None
        author_time = None

    if author_name is not None and author_time is not None:
        doc_name = doc.add_paragraph(author_name)
        doc_time = doc.add_paragraph(author_time.strftime("%Y/%m/%d %H:%M"))

        for run in doc_name.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(17 / 1.5)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for run in doc_time.runs:
            run.font.color.rgb = RGBColor(148, 153, 160)
            run.font.size = Pt(13 / 1.5)

    top_pic_url = list()
    if module_top is not None:
        module_top = module_top['module_top']['display']['album']['pics']
        for pic in module_top:
            top_pic_url.append(pic['url'])

    add_image(doc, get_img(cookie, top_pic_url, 'img'))

    if module_content is not None:
        module_content = module_content['module_content']['paragraphs']
        print(module_content)

    for item in module_content:
        if item['para_type'] == 1:
            add_text(doc, item['text'])
        elif item['para_type'] == 2:
            print('image')
        elif item['para_type'] == 3:
            print('i don\'t know')
        elif item['para_type'] == 4:
            print('code_text')
        elif item['para_type'] == 5:
            print('i don\'t know')
        elif item['para_type'] == 6:
            print('link_card')
        print(item)
        print('-*-'*100)

    print(title, author_name, author_time)

    doc.save(f'{doc_storage_location}/{document_name}')


if __name__ == '__main__':
    with open('../cookie.json', 'r') as f:
        cookie = json.load(f)['cookie']

    article_id = '318740526742916934'
    get_article(cookie, article_id,'./', 'document.doc')
